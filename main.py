import streamlit as st
import fitz  # PyMuPDF
import docx
from docx import Document
from deep_translator import GoogleTranslator
import os
import tempfile
import io

try:
    import pythoncom
    import win32com.client as win32
except ImportError:
    pythoncom = None
    win32 = None

import platform

# Função para extrair texto de PDFs, página por página
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        pages.append(text)
    return pages

# Função para extrair texto de arquivos .docx
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# Função para converter e extrair texto de arquivos .doc
def extract_text_from_doc(doc_path):
    if platform.system() == "Windows" and pythoncom and win32:
        try:
            pythoncom.CoInitialize()
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            temp_docx_path = doc_path + 'x'
            doc.SaveAs(temp_docx_path, FileFormat=16)  # 16 é o formato para .docx
            doc.Close()
            word.Quit()
            return extract_text_from_docx(temp_docx_path)
        except Exception as e:
            st.error(f"Erro ao converter arquivo .doc: {e}")
            return ""
    else:
        try:
            import subprocess
            temp_docx_path = doc_path + 'x'
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', os.path.dirname(doc_path)])
            return extract_text_from_docx(temp_docx_path)
        except Exception as e:
            st.error(f"Erro ao converter arquivo .doc no Linux: {e}")
            return ""

# Função para dividir o texto em partes menores
def split_text(text, max_length=5000):
    parts = []
    while len(text) > max_length:
        split_index = text[:max_length].rfind('\n')
        if split_index == -1:
            split_index = max_length
        parts.append(text[:split_index])
        text = text[split_index:]
    parts.append(text)
    return parts

# Função principal para processar documentos
def process_documents(file_paths, lingua_destino):
    translator = GoogleTranslator(source='auto', target=lingua_destino)
    translated_files = []

    for file_path, original_filename in file_paths:
        if file_path.endswith('.pdf'):
            pages = extract_text_from_pdf(file_path)
            file_type = 'pdf'
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)
            pages = [text]
            file_type = 'docx'
        elif file_path.endswith('.doc'):
            text = extract_text_from_doc(file_path)
            pages = [text]
            file_type = 'doc'
            if text == "":
                continue
        else:
            st.error(f"Formato não suportado: {file_path}")
            continue

        translated_pages = []
        for page in pages:
            parts = split_text(page)
            translated_text = ""
            for part in parts:
                translated_text += translator.translate(part) + "\n"
            translated_pages.append(translated_text.strip())

        translated_files.append((translated_pages, original_filename, file_type))

    return translated_files

# Função para criar um novo documento docx com o texto traduzido
def create_translated_docx(translated_pages, original_filename):
    doc = Document()
    for translated_page in translated_pages:
        for line in translated_page.split('\n'):
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    translated_filename = f"{os.path.splitext(original_filename)[0]}_traduzido.docx"
    return buffer, translated_filename

# Função para criar um novo documento PDF com o texto traduzido
def create_translated_pdf(translated_pages, original_filename):
    doc = fitz.open()
    for translated_page in translated_pages:
        parts = split_text(translated_page, max_length=1000)
        for part in parts:
            page = doc.new_page()
            page.insert_text((72, 72), part)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    translated_filename = f"{os.path.splitext(original_filename)[0]}_traduzido.pdf"
    return buffer, translated_filename

# Interface do Streamlit
st.title("Tradutor de Documentos")

st.markdown("""
### Anexe seus documentos para tradução
Suporte para arquivos: **.pdf, .doc, .docx**
""")

uploaded_files = st.file_uploader("Anexe arquivos de documentos", type=["pdf", "doc", "docx"], accept_multiple_files=True)

lingua_destino = st.selectbox("Selecione a língua de destino", ["pt", "en", "es", "fr", "de"])

formato_destino = st.selectbox("Selecione o formato do arquivo traduzido", ["PDF", "DOCX"])

if st.button("Traduzir"):
    if uploaded_files:
        st.session_state.tradução_em_andamento = True
        with st.spinner('Processando e traduzindo documentos...'):
            file_paths = []
            for uploaded_file in uploaded_files:
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
                    temp_file.write(uploaded_file.read())
                    file_paths.append((temp_file.name, uploaded_file.name))
            
            translated_files = process_documents(file_paths, lingua_destino)

            for translated_pages, original_filename, file_type in translated_files:
                if formato_destino == 'PDF':
                    buffer, translated_filename = create_translated_pdf(translated_pages, original_filename)
                elif formato_destino == 'DOCX':
                    buffer, translated_filename = create_translated_docx(translated_pages, original_filename)
                
                st.success(f'Tradução concluída para: {original_filename}')
                st.download_button(
                    label=f"Baixar {translated_filename}",
                    data=buffer,
                    file_name=translated_filename,
                    mime="application/pdf" if formato_destino == 'PDF' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
        st.session_state.tradução_em_andamento = False
    else:
        st.error("Por favor, anexe pelo menos um arquivo para traduzir.")
